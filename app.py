"""
TeachAI - AI Teacher Assistant Web App
Flask Backend - Main Application
"""

import os
import json
import uuid
import tempfile
from datetime import datetime
from flask import Flask, render_template, request, jsonify, send_file, redirect, url_for, flash
from werkzeug.utils import secure_filename
import sqlite3

app = Flask(__name__)
app.secret_key = "teachai_secret_key_2024"

# ─── CONFIG ────────────────────────────────────────────────────
UPLOAD_FOLDER = os.path.join(os.path.dirname(__file__), "uploads")
OUTPUT_FOLDER = os.path.join(os.path.dirname(__file__), "outputs")
DATABASE = os.path.join(os.path.dirname(__file__), "teachai.db")
ALLOWED_EXTENSIONS = {"xlsx", "xls", "png", "jpg", "jpeg", "pdf"}

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["OUTPUT_FOLDER"] = OUTPUT_FOLDER
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024  # 16MB max

# ─── DATABASE INIT ─────────────────────────────────────────────
def init_db():
    conn = sqlite3.connect(DATABASE)
    c = conn.cursor()
    c.execute("""
        CREATE TABLE IF NOT EXISTS students (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            roll_number TEXT,
            class_name TEXT,
            section TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)
    c.execute("""
        CREATE TABLE IF NOT EXISTS marks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            student_id INTEGER,
            subject TEXT NOT NULL,
            marks_obtained REAL,
            max_marks REAL DEFAULT 100,
            exam_type TEXT DEFAULT 'Unit Test',
            exam_date DATE,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (student_id) REFERENCES students(id)
        )
    """)
    c.execute("""
        CREATE TABLE IF NOT EXISTS history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            file_name TEXT NOT NULL,
            file_type TEXT NOT NULL,
            file_path TEXT NOT NULL,
            description TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)
    c.execute("""
        CREATE TABLE IF NOT EXISTS question_papers (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            school_name TEXT,
            subject TEXT,
            class_name TEXT,
            max_marks INTEGER,
            time_duration TEXT,
            difficulty TEXT,
            questions TEXT,
            file_path TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)
    conn.commit()
    conn.close()

def get_db():
    conn = sqlite3.connect(DATABASE)
    conn.row_factory = sqlite3.Row
    return conn

# ─── HELPERS ────────────────────────────────────────────────────
def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

def calculate_grade(percentage):
    if percentage >= 90: return "A+"
    elif percentage >= 80: return "A"
    elif percentage >= 70: return "B+"
    elif percentage >= 60: return "B"
    elif percentage >= 50: return "C"
    elif percentage >= 40: return "D"
    else: return "F"

def fuzzy_match(name1, name2):
    """Simple fuzzy matching for student names"""
    n1 = name1.lower().strip()
    n2 = name2.lower().strip()
    if n1 == n2: return 100
    if n1 in n2 or n2 in n1: return 85
    # Character overlap ratio
    set1, set2 = set(n1.split()), set(n2.split())
    if set1 & set2:
        ratio = len(set1 & set2) / max(len(set1), len(set2))
        return int(ratio * 80)
    return 0

def save_to_history(filename, file_type, file_path, description=""):
    conn = get_db()
    conn.execute(
        "INSERT INTO history (file_name, file_type, file_path, description) VALUES (?, ?, ?, ?)",
        (filename, file_type, file_path, description)
    )
    conn.commit()
    conn.close()

# ─── ROUTES – PAGES ─────────────────────────────────────────────
@app.route("/")
def dashboard():
    conn = get_db()
    student_count = conn.execute("SELECT COUNT(*) FROM students").fetchone()[0]
    marks_count = conn.execute("SELECT COUNT(*) FROM marks").fetchone()[0]
    paper_count = conn.execute("SELECT COUNT(*) FROM question_papers").fetchone()[0]
    history_count = conn.execute("SELECT COUNT(*) FROM history").fetchone()[0]
    recent_history = conn.execute(
        "SELECT * FROM history ORDER BY created_at DESC LIMIT 5"
    ).fetchall()
    conn.close()
    return render_template("dashboard.html",
        student_count=student_count,
        marks_count=marks_count,
        paper_count=paper_count,
        history_count=history_count,
        recent_history=recent_history
    )

@app.route("/marks")
def marks_page():
    conn = get_db()
    students = conn.execute("SELECT * FROM students ORDER BY name").fetchall()
    subjects = conn.execute("SELECT DISTINCT subject FROM marks").fetchall()
    conn.close()
    return render_template("marks.html", students=students, subjects=subjects)

@app.route("/paper-generator")
def paper_generator():
    return render_template("paper_generator.html")

@app.route("/history")
def history_page():
    conn = get_db()
    files = conn.execute(
        "SELECT * FROM history ORDER BY created_at DESC"
    ).fetchall()
    conn.close()
    return render_template("history.html", files=files)

# ─── API – STUDENTS ─────────────────────────────────────────────
@app.route("/api/students", methods=["GET"])
def get_students():
    conn = get_db()
    students = conn.execute("SELECT * FROM students ORDER BY name").fetchall()
    conn.close()
    return jsonify([dict(s) for s in students])

@app.route("/api/students", methods=["POST"])
def add_student():
    data = request.json
    if not data or not data.get("name"):
        return jsonify({"error": "Name is required"}), 400
    conn = get_db()
    conn.execute(
        "INSERT INTO students (name, roll_number, class_name, section) VALUES (?, ?, ?, ?)",
        (data["name"], data.get("roll_number", ""), data.get("class_name", ""), data.get("section", ""))
    )
    conn.commit()
    conn.close()
    return jsonify({"success": True, "message": "Student added successfully"})

@app.route("/api/upload-students", methods=["POST"])
def upload_students():
    try:
        import pandas as pd
    except ImportError:
        return jsonify({"error": "pandas not installed. Run: pip install pandas openpyxl"}), 500

    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    file = request.files["file"]
    if not file or not allowed_file(file.filename):
        return jsonify({"error": "Invalid file type. Use .xlsx or .xls"}), 400

    filename = secure_filename(file.filename)
    filepath = os.path.join(UPLOAD_FOLDER, f"{uuid.uuid4()}_{filename}")
    file.save(filepath)

    try:
        df = pd.read_excel(filepath)
        df.columns = [c.lower().strip() for c in df.columns]

        name_col = next((c for c in df.columns if "name" in c), None)
        roll_col = next((c for c in df.columns if "roll" in c), None)
        class_col = next((c for c in df.columns if "class" in c), None)

        if not name_col:
            return jsonify({"error": "Excel must have a 'Name' column"}), 400

        conn = get_db()
        added = 0
        for _, row in df.iterrows():
            name = str(row[name_col]).strip()
            if not name or name == "nan": continue
            roll = str(row[roll_col]).strip() if roll_col else ""
            cls = str(row[class_col]).strip() if class_col else ""
            conn.execute(
                "INSERT INTO students (name, roll_number, class_name) VALUES (?, ?, ?)",
                (name, roll, cls)
            )
            added += 1
        conn.commit()
        conn.close()
        os.remove(filepath)
        return jsonify({"success": True, "message": f"{added} students imported successfully"})
    except Exception as e:
        return jsonify({"error": f"Failed to parse Excel: {str(e)}"}), 500

# ─── API – MARKS ─────────────────────────────────────────────────
@app.route("/api/marks", methods=["POST"])
def add_marks():
    data = request.json
    required = ["student_id", "subject", "marks_obtained"]
    if not all(k in data for k in required):
        return jsonify({"error": "Missing required fields"}), 400

    max_marks = float(data.get("max_marks", 100))
    marks = float(data["marks_obtained"])
    if marks < 0 or marks > max_marks:
        return jsonify({"error": f"Marks must be between 0 and {max_marks}"}), 400

    conn = get_db()
    conn.execute(
        "INSERT INTO marks (student_id, subject, marks_obtained, max_marks, exam_type, exam_date) VALUES (?, ?, ?, ?, ?, ?)",
        (data["student_id"], data["subject"], marks, max_marks,
         data.get("exam_type", "Unit Test"), data.get("exam_date", datetime.now().date()))
    )
    conn.commit()
    conn.close()
    return jsonify({"success": True, "message": "Marks saved"})

@app.route("/api/marks/voice", methods=["POST"])
def process_voice_marks():
    """Parse voice input: 'Rahul 78 Aman 85 Priya 92'"""
    data = request.json
    text = data.get("text", "").strip()
    subject = data.get("subject", "General")
    if not text:
        return jsonify({"error": "No voice text provided"}), 400

    tokens = text.replace(",", " ").split()
    results = []
    i = 0
    while i < len(tokens) - 1:
        try:
            marks = float(tokens[i + 1])
            name = tokens[i]
            results.append({"name": name, "marks": marks})
            i += 2
        except ValueError:
            i += 1

    conn = get_db()
    students = conn.execute("SELECT * FROM students").fetchall()
    matched = []
    unmatched = []

    for entry in results:
        best_match = None
        best_score = 0
        for s in students:
            score = fuzzy_match(entry["name"], s["name"])
            if score > best_score:
                best_score = score
                best_match = s
        if best_match and best_score >= 60:
            conn.execute(
                "INSERT INTO marks (student_id, subject, marks_obtained) VALUES (?, ?, ?)",
                (best_match["id"], subject, entry["marks"])
            )
            matched.append({"name": best_match["name"], "marks": entry["marks"], "confidence": best_score})
        else:
            unmatched.append(entry["name"])

    conn.commit()
    conn.close()
    return jsonify({
        "success": True,
        "matched": matched,
        "unmatched": unmatched,
        "message": f"{len(matched)} marks saved, {len(unmatched)} not matched"
    })

@app.route("/api/marks/upload", methods=["POST"])
def upload_marks():
    try:
        import pandas as pd
    except ImportError:
        return jsonify({"error": "pandas not installed"}), 500

    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]
    subject = request.form.get("subject", "General")

    filename = secure_filename(file.filename)
    ext = filename.rsplit(".", 1)[-1].lower()

    if ext in ["png", "jpg", "jpeg"]:
        return process_ocr_marks(file, subject)

    filepath = os.path.join(UPLOAD_FOLDER, f"{uuid.uuid4()}_{filename}")
    file.save(filepath)

    try:
        df = pd.read_excel(filepath)
        df.columns = [str(c).lower().strip() for c in df.columns]
        name_col = next((c for c in df.columns if "name" in c), None)
        marks_col = next((c for c in df.columns if "mark" in c or "score" in c), None)

        if not name_col or not marks_col:
            return jsonify({"error": "Excel must have 'Name' and 'Marks' columns"}), 400

        conn = get_db()
        students = conn.execute("SELECT * FROM students").fetchall()
        matched, errors = [], []

        for _, row in df.iterrows():
            name = str(row[name_col]).strip()
            if not name or name == "nan": continue
            try:
                marks = float(row[marks_col])
            except:
                errors.append(f"Invalid marks for {name}")
                continue
            best, score = None, 0
            for s in students:
                sc = fuzzy_match(name, s["name"])
                if sc > score: score, best = sc, s
            if best and score >= 60:
                conn.execute(
                    "INSERT INTO marks (student_id, subject, marks_obtained) VALUES (?, ?, ?)",
                    (best["id"], subject, marks)
                )
                matched.append(name)
            else:
                errors.append(f"No match found for: {name}")

        conn.commit()
        conn.close()
        os.remove(filepath)
        return jsonify({"success": True, "matched": len(matched), "errors": errors})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

def process_ocr_marks(file, subject):
    """OCR from image using Tesseract"""
    try:
        import pytesseract
        from PIL import Image
        import io

        img = Image.open(io.BytesIO(file.read()))
        text = pytesseract.image_to_string(img)
        # Reuse voice parser logic
        from flask import current_app
        with current_app.test_request_context(
            json={"text": text, "subject": subject}
        ):
            return process_voice_marks()
    except ImportError:
        return jsonify({"error": "Tesseract/PIL not installed. Install: pip install pytesseract Pillow"}), 500
    except Exception as e:
        return jsonify({"error": f"OCR failed: {str(e)}"}), 500

@app.route("/api/marks/report", methods=["GET"])
def get_marks_report():
    conn = get_db()
    data = conn.execute("""
        SELECT s.name, s.roll_number, s.class_name, m.subject,
               m.marks_obtained, m.max_marks, m.exam_type, m.exam_date
        FROM marks m JOIN students s ON m.student_id = s.id
        ORDER BY s.name, m.subject
    """).fetchall()
    conn.close()

    report = {}
    for row in data:
        name = row["name"]
        if name not in report:
            report[name] = {"name": name, "roll": row["roll_number"],
                            "class": row["class_name"], "subjects": {}, "total": 0, "max_total": 0}
        subj = row["subject"]
        marks = row["marks_obtained"] or 0
        max_m = row["max_marks"] or 100
        pct = round((marks / max_m) * 100, 1)
        report[name]["subjects"][subj] = {
            "marks": marks, "max": max_m,
            "percentage": pct, "grade": calculate_grade(pct)
        }
        report[name]["total"] += marks
        report[name]["max_total"] += max_m

    for name in report:
        if report[name]["max_total"] > 0:
            overall_pct = round((report[name]["total"] / report[name]["max_total"]) * 100, 1)
            report[name]["overall_percentage"] = overall_pct
            report[name]["overall_grade"] = calculate_grade(overall_pct)

    return jsonify(list(report.values()))

@app.route("/api/marks/export", methods=["GET"])
def export_marks_excel():
    try:
        import pandas as pd
    except ImportError:
        return jsonify({"error": "pandas not installed"}), 500

    conn = get_db()
    data = conn.execute("""
        SELECT s.name, s.roll_number, s.class_name, m.subject,
               m.marks_obtained, m.max_marks, m.exam_type
        FROM marks m JOIN students s ON m.student_id = s.id
        ORDER BY s.name, m.subject
    """).fetchall()
    conn.close()

    rows = []
    for row in data:
        marks = row["marks_obtained"] or 0
        max_m = row["max_marks"] or 100
        pct = round((marks / max_m) * 100, 1)
        rows.append({
            "Name": row["name"],
            "Roll No": row["roll_number"],
            "Class": row["class_name"],
            "Subject": row["subject"],
            "Marks": marks,
            "Max Marks": max_m,
            "Percentage": pct,
            "Grade": calculate_grade(pct),
            "Exam Type": row["exam_type"]
        })

    df = pd.DataFrame(rows)
    filename = f"marks_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    filepath = os.path.join(OUTPUT_FOLDER, filename)
    df.to_excel(filepath, index=False)
    save_to_history(filename, "Excel", filepath, "Marks Report")
    return send_file(filepath, as_attachment=True, download_name=filename)

# ─── API – PAPER GENERATOR ──────────────────────────────────────
@app.route("/api/generate-paper", methods=["POST"])
def generate_paper():
    data = request.json
    school = data.get("school_name", "School Name")
    subject = data.get("subject", "General")
    class_name = data.get("class_name", "10")
    max_marks = data.get("max_marks", 80)
    duration = data.get("duration", "3 Hours")
    difficulty = data.get("difficulty", "Medium")
    questions_raw = data.get("questions", "")
    output_format = data.get("format", "docx")

    # Parse questions
    questions = []
    if isinstance(questions_raw, str):
        lines = [l.strip() for l in questions_raw.strip().split("\n") if l.strip()]
        for line in lines:
            questions.append(line)
    elif isinstance(questions_raw, list):
        questions = questions_raw

    # Auto-generate placeholder questions if none provided
    if not questions:
        questions = generate_ai_questions(subject, class_name, difficulty)

    # Distribute into sections
    sec_a = questions[:5] if len(questions) >= 5 else questions
    sec_b = questions[5:10] if len(questions) >= 10 else questions[len(sec_a):]
    sec_c = questions[10:] if len(questions) > 10 else []

    paper_data = {
        "school": school, "subject": subject, "class": class_name,
        "max_marks": max_marks, "duration": duration, "difficulty": difficulty,
        "date": datetime.now().strftime("%d/%m/%Y"),
        "sec_a": sec_a, "sec_b": sec_b, "sec_c": sec_c
    }

    if output_format == "docx":
        return generate_docx_paper(paper_data)
    else:
        return generate_pdf_paper(paper_data)

def generate_ai_questions(subject, class_name, difficulty):
    """AI-suggested placeholder questions based on subject"""
    question_bank = {
        "Math": [
            f"Solve: Find the value of x in the equation 2x + 5 = 15",
            f"Find the area of a circle with radius 7 cm",
            f"Factorize: x² + 5x + 6",
            f"A train travels 360 km in 4 hours. Find its speed.",
            f"Find the HCF of 36 and 48",
            f"Prove that √2 is irrational",
            f"Find the roots of the quadratic equation x² - 7x + 12 = 0",
            f"In a triangle, angles are in ratio 1:2:3. Find each angle.",
        ],
        "Science": [
            f"What is Newton's Second Law of Motion? Give an example.",
            f"Explain the process of photosynthesis",
            f"What is the difference between physical and chemical change?",
            f"Name the organs of the human digestive system",
            f"What is Ohm's Law? Write its formula.",
            f"Explain the water cycle with a diagram",
            f"What are the properties of acids and bases?",
        ],
        "English": [
            f"Write a paragraph on 'Importance of Trees'",
            f"Correct the following sentences: (a) She go to school daily.",
            f"Write the opposite of the following words: happy, dark, strong",
            f"Read the passage and answer the questions",
            f"Write a letter to your friend about your summer vacation",
        ],
        "Hindi": [
            f"निम्नलिखित शब्दों के अर्थ लिखिए: (क) विद्यालय (ख) आकाश",
            f"किसी एक विषय पर अनुच्छेद लिखिए: 'मेरा प्रिय खेल'",
            f"वाक्य में रेखांकित शब्द का पद परिचय दीजिए",
            f"संधि विच्छेद कीजिए: विद्यालय, रामायण",
        ]
    }
    base = question_bank.get(subject, [
        "Define and explain the main concept",
        "Give three examples from daily life",
        "Compare and contrast the given terms",
        "Describe the process in detail",
        "Solve the given problem step by step",
        "Explain with a diagram",
        "What are the main differences between A and B?",
        "Write short notes on any two",
    ])
    if difficulty == "Hard":
        base = [f"[Hard] {q}" for q in base]
    elif difficulty == "Easy":
        base = [f"[Easy] {q}" for q in base]
    return base

def generate_docx_paper(paper_data):
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return jsonify({"error": "python-docx not installed. Run: pip install python-docx"}), 500

    doc = Document()

    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(paper_data["school"].upper())
    run.bold = True
    run.font.size = Pt(16)

    sub_header = doc.add_paragraph()
    sub_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_header.add_run(f"Subject: {paper_data['subject']}  |  Class: {paper_data['class']}  |  Date: {paper_data['date']}")

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"Time: {paper_data['duration']}  |  Max Marks: {paper_data['max_marks']}  |  Difficulty: {paper_data['difficulty']}")

    doc.add_paragraph("─" * 60)
    doc.add_paragraph("General Instructions: All questions are compulsory. Write legibly.")
    doc.add_paragraph("─" * 60)

    def add_section(title, questions, marks_each):
        if not questions: return
        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(12)
        for i, q in enumerate(questions, 1):
            p = doc.add_paragraph()
            p.add_run(f"Q{i}. {q}")
            p.add_run(f"  [{marks_each} marks]").bold = True

    add_section("SECTION A – Short Answer Questions (2 marks each)", paper_data["sec_a"], 2)
    add_section("SECTION B – Long Answer Questions (5 marks each)", paper_data["sec_b"], 5)
    if paper_data["sec_c"]:
        add_section("SECTION C – Optional Questions (Any 2 of the following)", paper_data["sec_c"], 8)

    filename = f"question_paper_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(OUTPUT_FOLDER, filename)
    doc.save(filepath)

    # Save to DB
    conn = get_db()
    conn.execute(
        "INSERT INTO question_papers (school_name, subject, class_name, max_marks, time_duration, difficulty, file_path) VALUES (?,?,?,?,?,?,?)",
        (paper_data["school"], paper_data["subject"], paper_data["class"],
         paper_data["max_marks"], paper_data["duration"], paper_data["difficulty"], filepath)
    )
    conn.commit()
    conn.close()
    save_to_history(filename, "Word Document", filepath, f"{paper_data['subject']} Question Paper")

    return send_file(filepath, as_attachment=True, download_name=filename,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

def generate_pdf_paper(paper_data):
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.lib import colors
    except ImportError:
        return jsonify({"error": "reportlab not installed. Run: pip install reportlab"}), 500

    filename = f"question_paper_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    filepath = os.path.join(OUTPUT_FOLDER, filename)

    doc = SimpleDocTemplate(filepath, pagesize=A4,
                            rightMargin=2*cm, leftMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    story = []

    title_style = ParagraphStyle("title", parent=styles["Heading1"],
                                  fontSize=18, alignment=TA_CENTER, spaceAfter=6)
    sub_style = ParagraphStyle("sub", parent=styles["Normal"],
                                fontSize=11, alignment=TA_CENTER, spaceAfter=4)
    section_style = ParagraphStyle("section", parent=styles["Heading2"],
                                    fontSize=13, spaceAfter=6)
    q_style = ParagraphStyle("q", parent=styles["Normal"],
                               fontSize=11, leftIndent=10, spaceAfter=8)

    story.append(Paragraph(paper_data["school"].upper(), title_style))
    story.append(Paragraph(f"Subject: {paper_data['subject']} &nbsp;&nbsp; Class: {paper_data['class']} &nbsp;&nbsp; Date: {paper_data['date']}", sub_style))
    story.append(Paragraph(f"Time: {paper_data['duration']} &nbsp;&nbsp; Max Marks: {paper_data['max_marks']} &nbsp;&nbsp; Difficulty: {paper_data['difficulty']}", sub_style))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.black))
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph("General Instructions: All questions are compulsory. Write legibly.", styles["Italic"]))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.grey))
    story.append(Spacer(1, 0.5*cm))

    def add_pdf_section(title, questions, marks_each):
        if not questions: return
        story.append(Paragraph(title, section_style))
        for i, q in enumerate(questions, 1):
            story.append(Paragraph(f"<b>Q{i}.</b> {q} &nbsp;&nbsp; <i>[{marks_each} marks]</i>", q_style))
        story.append(Spacer(1, 0.4*cm))

    add_pdf_section("SECTION A – Short Answer Questions (2 marks each)", paper_data["sec_a"], 2)
    add_pdf_section("SECTION B – Long Answer Questions (5 marks each)", paper_data["sec_b"], 5)
    if paper_data["sec_c"]:
        add_pdf_section("SECTION C – Optional Questions", paper_data["sec_c"], 8)

    doc.build(story)

    save_to_history(filename, "PDF", filepath, f"{paper_data['subject']} Question Paper")
    return send_file(filepath, as_attachment=True, download_name=filename, mimetype="application/pdf")

# ─── API – HISTORY ───────────────────────────────────────────────
@app.route("/api/history", methods=["GET"])
def get_history():
    conn = get_db()
    files = conn.execute("SELECT * FROM history ORDER BY created_at DESC LIMIT 50").fetchall()
    conn.close()
    return jsonify([dict(f) for f in files])

@app.route("/api/download/<int:file_id>")
def download_file(file_id):
    conn = get_db()
    file = conn.execute("SELECT * FROM history WHERE id = ?", (file_id,)).fetchone()
    conn.close()
    if not file or not os.path.exists(file["file_path"]):
        return jsonify({"error": "File not found"}), 404
    return send_file(file["file_path"], as_attachment=True, download_name=file["file_name"])

@app.route("/api/history/<int:file_id>", methods=["DELETE"])
def delete_history(file_id):
    conn = get_db()
    file = conn.execute("SELECT * FROM history WHERE id = ?", (file_id,)).fetchone()
    if file:
        try: os.remove(file["file_path"])
        except: pass
        conn.execute("DELETE FROM history WHERE id = ?", (file_id,))
        conn.commit()
    conn.close()
    return jsonify({"success": True})

@app.route("/api/stats")
def get_stats():
    conn = get_db()
    stats = {
        "students": conn.execute("SELECT COUNT(*) FROM students").fetchone()[0],
        "marks_entries": conn.execute("SELECT COUNT(*) FROM marks").fetchone()[0],
        "papers_generated": conn.execute("SELECT COUNT(*) FROM question_papers").fetchone()[0],
        "files_generated": conn.execute("SELECT COUNT(*) FROM history").fetchone()[0],
        "subjects": [r[0] for r in conn.execute("SELECT DISTINCT subject FROM marks").fetchall()]
    }
    conn.close()
    return jsonify(stats)

# ─── RUN ─────────────────────────────────────────────────────────
if __name__ == "__main__":
    init_db()
    print("=" * 50)
    print("  TeachAI - AI Teacher Assistant")
    print("  Running at: http://localhost:5000")
    print("=" * 50)
    app.run(debug=True, host="0.0.0.0", port=5000)
